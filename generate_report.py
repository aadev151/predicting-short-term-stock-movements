from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import copy
import pandas as pd

OUTPUT_PATH = "draft_report.docx"

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── Default body font ────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(12)
style.paragraph_format.space_after  = Pt(6)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = Pt(14)   # ~single spacing

# ── Heading styles ───────────────────────────────────────────────────────────
h1 = doc.styles['Heading 1']
h1.font.name  = 'Calibri'
h1.font.size  = Pt(16)
h1.font.bold  = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.space_before = Pt(14)
h1.paragraph_format.space_after  = Pt(4)

h2 = doc.styles['Heading 2']
h2.font.name  = 'Calibri'
h2.font.size  = Pt(13)
h2.font.bold  = True
h2.font.italic = False
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before = Pt(10)
h2.paragraph_format.space_after  = Pt(2)

# ── Helpers ──────────────────────────────────────────────────────────────────

def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    return p


def set_cell_border(cell, **kwargs):
    """Apply borders to a table cell via XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), kwargs.get('space', '0'))
        tag.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def make_table(headers, rows, col_widths=None):
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold      = True
        run.font.size = Pt(11)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(hdr_cells[i])

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(11)
            cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_border(cells[c_idx])

    # Optional column widths
    if col_widths:
        for row in table.rows:
            for c_idx, w in enumerate(col_widths):
                row.cells[c_idx].width = Inches(w)

    doc.add_paragraph()   # spacing after table
    return table


def add_numbered_list(items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style='List Number')
        p.style = doc.styles['Normal']
        p.paragraph_format.left_indent  = Inches(0.25)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(f"{i}. {item}")
        run.font.size = Pt(12)


def add_bullet(text, indent=0.25):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"\u2022 {text}")
    run.font.size = Pt(12)
    return p


def add_checklist_item(text, checked=True, indent=0.25):
    symbol = '\u2611' if checked else '\u2610'   # ☑ / ☐
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"{symbol} {text}")
    run.font.size = Pt(12)
    return p

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE BLOCK
# ════════════════════════════════════════════════════════════════════════════

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run(
    "Predicting Short-Term Stock Price Direction Using\n"
    "Technical Indicators and News Sentiment"
)
title_run.bold      = True
title_run.font.name = 'Calibri'
title_run.font.size = Pt(20)
title_p.paragraph_format.space_after = Pt(10)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run(
    "Draft Report \u2014 Team members: Alex, Seb, Umesh, Naziia"
)
sub_run.italic    = True
sub_run.font.name = 'Calibri'
sub_run.font.size = Pt(13)
sub_p.paragraph_format.space_after = Pt(20)

doc.add_paragraph()   # blank spacer

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Introduction', level=1)

add_body(
    "Stock price prediction is one of the most studied problems in quantitative finance, "
    "yet it remains notoriously difficult. Efficient Market Hypothesis (EMH) suggests that "
    "publicly available information is already priced in, leaving little exploitable signal. "
    "This project tests a narrower, more concrete version of the question: Can a combination "
    "of technical indicators, macroeconomic signals, and news sentiment predict whether an "
    "individual S\u0026P 500 stock will close higher or lower the following trading day?"
)

add_body(
    "We frame this as a binary classification task (Up\u00a0=\u00a01, Down/Flat\u00a0=\u00a00) "
    "and evaluate models against four baselines designed to represent progressively smarter "
    "zero-effort strategies."
)

# ── Subsection: Dataset ───────────────────────────────────────────────────
doc.add_heading('Dataset', level=2)

add_body("Our dataset was assembled from six sources:")

dataset_sources = [
    (
        "S\u0026P 500 price data",
        " \u2014 daily OHLCV and adjusted close for current S\u0026P 500 constituents, "
        "downloaded from Yahoo Finance via yfinance. Covers 2016-02-10 to 2026-02-09, "
        "yielding approximately 1.26 million stock-day rows across 503 tickers. Note: we "
        "intentionally use the current constituent list for all historical dates, introducing "
        "survivorship bias."
    ),
    (
        "Stock rolling / window-based features (JP)",
        " \u2014 technical factors derived from historical OHLCV data, including multi-horizon "
        "returns and momentum, rolling volatility and mean returns, ATR/range statistics, "
        "volume shock measures, RSI, MACD, Bollinger Bands, and 60-day rolling beta/correlation "
        "versus the S\u0026P 500."
    ),
    (
        "Company metadata and index level (Umesh)",
        " \u2014 sector, industry, market cap, revenue growth, S\u0026P 500 index weight, "
        "and the daily S\u0026P 500 index price."
    ),
    (
        "NYSE OHLC data (Naziia)",
        " \u2014 alternative open/high/low/close series for a subset of tickers."
    ),
    (
        "Market regime features (Seb)",
        " \u2014 daily VIX (standardised), yield spread, and a Gaussian Mixture Model (GMM) "
        "regime label with three states (0\u00a0= calm, 1\u00a0= normal, 2\u00a0= stressful)."
    ),
    (
        "News sentiment (Alex)",
        " \u2014 FinBERT sentiment scores aggregated by (date, ticker) from the Kaggle "
        "\u201cMassive Stock News\u201d dataset (analyst headlines, 2016\u20132020). Features "
        "include news_count, sentiment_mean, sentiment_ratio, and a binary has_news flag."
    ),
]

for i, (bold_part, rest) in enumerate(dataset_sources, 1):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    run_num  = p.add_run(f"{i}. ")
    run_bold = p.add_run(bold_part)
    run_bold.bold = True
    run_bold.font.size = Pt(12)
    run_num.font.size  = Pt(12)
    run_rest = p.add_run(rest)
    run_rest.font.size = Pt(12)

doc.add_paragraph()

add_body(
    "The final merged dataset contains 43 columns. Missingness is substantial for several "
    "data streams: NYSE prices (~85\u0025 missing), market regime features (~84\u0025 "
    "missing), and news sentiment (~94\u0025 missing). Missing values for modelling are "
    "handled via median imputation (VIX, Yield Spread, Regime) and zero-fill (sentiment, "
    "revenue growth), so all rows are retained after feature engineering."
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — METHODOLOGY
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Methodology', level=1)

# ── Feature Engineering ──────────────────────────────────────────────────
doc.add_heading('Feature Engineering', level=2)

add_body(
    "We derived 18 features from the merged dataset. After dropping rows with remaining NaN "
    "values, 1,219,484 rows are used for modelling. The feature groups are:"
)

feature_groups = [
    ("Price momentum \u0026 lagged returns", ": lag_return_1, lag_return_2, lag_return_5, price_to_ma20, rolling_std_20"),
    ("Intraday structure", ": hl_range (high\u2013low / close), oc_gap (overnight gap), vol_norm (volume vs 20-day average)"),
    ("Macro \u0026 regime", ": VIX, Yield_Spread, Regime_GMM"),
    ("Sentiment \u0026 fundamentals", ": sentiment_mean, sentiment_ratio, has_news, log_marketcap, Revenuegrowth, Weight, Sector_encoded"),
]

for bold_label, rest in feature_groups:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("\u2022 ").font.size = Pt(12)
    rb = p.add_run(bold_label)
    rb.bold = True
    rb.font.size = Pt(12)
    rr = p.add_run(rest)
    rr.font.size = Pt(12)

doc.add_paragraph()

add_body(
    "The target variable is target\u00a0=\u00a01 if next-day return\u00a0>\u00a00, else\u00a00. "
    "No future information is used in any feature; all values are observable at market close "
    "on day\u00a0T to predict day\u00a0T+1."
)

# ── Train / Test Split ────────────────────────────────────────────────────
doc.add_heading('Train / Test Split', level=2)

add_body(
    "We use a time-based split at 2023-01-01, with no shuffling, to prevent leakage of "
    "future information into training. This yields:"
)

for item in [
    "Train: 830,418 rows (68.3\u0025), spanning 2016 through end of 2022.",
    "Test: 389,066 rows (31.7\u0025), spanning 2023 through early 2026.",
]:
    add_bullet(item)

add_body(
    "The test set deliberately includes distinct market environments (post-COVID recovery, "
    "2024 election period, and 2025), providing a stress test of out-of-sample generalization."
)

# ── Baselines ─────────────────────────────────────────────────────────────
doc.add_heading('Baselines', level=2)

add_body(
    "Before evaluating any ML model, we established four zero-parameter baselines:"
)

baselines = [
    ("Random Guess", " \u2014 Coin flip; absolute floor."),
    ("Always Up", " \u2014 Exploit positive long-run market drift."),
    ("Follow Yesterday", " \u2014 Predict tomorrow = direction of today\u2019s return (individual momentum)."),
    ("Follow Market", " \u2014 Predict tomorrow = direction of today\u2019s S\u0026P\u00a0500 return."),
]

for bold_label, rest in baselines:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("\u2022 ").font.size = Pt(12)
    rb = p.add_run(bold_label)
    rb.bold = True
    rb.font.size = Pt(12)
    rr = p.add_run(rest)
    rr.font.size = Pt(12)

doc.add_paragraph()

add_body(
    'The "Follow Market" baseline is especially important because individual stock returns are '
    "largely driven by market-wide beta. Beating it demonstrates that a model has learned "
    "something beyond simple co-movement."
)

# ── Models ────────────────────────────────────────────────────────────────
doc.add_heading('Models', level=2)

add_body(
    "We trained and evaluated five classification models, all using the identical 18 features "
    "and the same time-split:"
)

p_dt = doc.add_paragraph(style='Normal')
p_dt.paragraph_format.left_indent = Inches(0.25)
p_dt.paragraph_format.space_after = Pt(4)
r = p_dt.add_run("Decision Tree")
r.bold = True
r.font.size = Pt(12)
p_dt.add_run(
    " \u2014 trained with GridSearchCV over max_depth in [2,\u200939] using 5-fold TimeSeriesSplit. "
    "The best depth (19) was selected on mean validation accuracy."
).font.size = Pt(12)

p_lr = doc.add_paragraph(style='Normal')
p_lr.paragraph_format.left_indent = Inches(0.25)
p_lr.paragraph_format.space_after = Pt(4)
r2 = p_lr.add_run("Logistic Regression")
r2.bold = True
r2.font.size = Pt(12)
p_lr.add_run(
    " \u2014 features were standardised with StandardScaler fitted on the training set only "
    "(no leakage). Regularisation strength was selected via LogisticRegressionCV with "
    "100 candidate C values (L2 penalty)."
).font.size = Pt(12)

p_rf = doc.add_paragraph(style='Normal')
p_rf.paragraph_format.left_indent = Inches(0.25)
p_rf.paragraph_format.space_after = Pt(4)
r3 = p_rf.add_run("Random Forest")
r3.bold = True
r3.font.size = Pt(12)
p_rf.add_run(
    " \u2014 300 trees with max_depth=8, min_samples_leaf=200, and max_features='sqrt' "
    "to reduce overfitting on the large panel dataset."
).font.size = Pt(12)

p_xgb = doc.add_paragraph(style='Normal')
p_xgb.paragraph_format.left_indent = Inches(0.25)
p_xgb.paragraph_format.space_after = Pt(4)
r4 = p_xgb.add_run("XGBoost")
r4.bold = True
r4.font.size = Pt(12)
p_xgb.add_run(
    " \u2014 gradient-boosted trees with regularisation (subsample/colsample, L1/L2, "
    "min_child_weight) trained on the same features and split."
).font.size = Pt(12)

p_hgb = doc.add_paragraph(style='Normal')
p_hgb.paragraph_format.left_indent = Inches(0.25)
p_hgb.paragraph_format.space_after = Pt(4)
r5 = p_hgb.add_run("HistGradientBoosting")
r5.bold = True
r5.font.size = Pt(12)
p_hgb.add_run(
    " \u2014 histogram-based gradient boosting (learning_rate=0.05, max_iter=300, "
    "max_depth=8, min_samples_leaf=200) as an additional non-linear baseline."
).font.size = Pt(12)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — RESULTS AND DISCUSSION
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Results and Discussion', level=1)

# ── Baseline Results ──────────────────────────────────────────────────────
doc.add_heading('Baseline Results', level=2)

p_cap1 = doc.add_paragraph(style='Normal')
p_cap1.paragraph_format.space_after = Pt(4)
r = p_cap1.add_run("Table 1: ")
r.bold = True
r.font.size = Pt(12)
p_cap1.add_run("Baseline Accuracy on Test Set (\u2265 2023-01-01)").font.size = Pt(12)

make_table(
    headers=["Baseline Strategy", "Test Accuracy"],
    rows=[
        ["Random Guess",    "0.5003"],
        ["Always Up",       "0.5224 (best)"],
        ["Follow Yesterday","0.4961"],
        ["Follow Market",   "0.4881"],
    ],
    col_widths=[3.0, 2.5],
)

add_body(
    'The "Always Up" baseline is the strongest, reflecting a slight positive skew in the test '
    "period (52.24\u0025 of days are up days). Notably, both \u201cFollow Yesterday\u201d and "
    "\u201cFollow Market\u201d perform below random chance on this test set, suggesting that "
    "short-term momentum is contrarian rather than persistent over 2023\u20132026."
)

# ── Model Results ─────────────────────────────────────────────────────────
doc.add_heading('Model Results', level=2)

p_cap2 = doc.add_paragraph(style='Normal')
p_cap2.paragraph_format.space_after = Pt(4)
r = p_cap2.add_run("Table 2: ")
r.bold = True
r.font.size = Pt(12)
p_cap2.add_run("Model Performance vs. Best Baseline (same features and split)").font.size = Pt(12)

make_table(
    headers=["Model", "Test Accuracy", "Balanced Accuracy", "\u0394 vs. Baseline"],
    rows=[
        ["Logistic Regression",      "0.5182", "0.5000", "\u22120.0042"],
        ["HistGradientBoosting",     "0.5180", "0.5010", "\u22120.0044"],
        ["Random Forest",            "0.5176", "0.4992", "\u22120.0048"],
        ["XGBoost",                  "0.5173", "0.5006", "\u22120.0051"],
        ["Decision Tree (depth 19)", "0.5170", "0.4991", "\u22120.0054"],
    ],
    col_widths=[2.5, 1.4, 1.6, 1.3],
)

add_body(
    'None of the five models beat the "Always Up" baseline on the held-out test set. '
    "Logistic Regression has the best raw accuracy (0.5182), while HistGradientBoosting has "
    "the best balanced accuracy (0.5010)."
)

add_body(
    "All models exhibit a similar failure mode: they strongly favour \u201cUp\u201d predictions. "
    "Down-class recall remains low (roughly 0.09\u20130.13), while Up-class recall remains high "
    "(roughly 0.87\u20130.91), keeping overall behaviour close to the Always-Up baseline."
)

# ── Per-Sector Results ───────────────────────────────────────────────────
doc.add_heading('Per-Sector Results', level=2)

sector_best_path = Path(__file__).resolve().parent / "data" / "model_best_by_sector.csv"
sector_model_path = Path(__file__).resolve().parent / "data" / "model_sector_results.csv"

if sector_best_path.exists() and sector_model_path.exists():
    sector_best = pd.read_csv(sector_best_path)
    sector_model = pd.read_csv(sector_model_path)

    sector_best = sector_best.sort_values("Sector").reset_index(drop=True)
    n_sectors = len(sector_best)
    n_positive = int((sector_best["Delta_vs_Sector_Baseline"] > 0).sum())

    mean_delta = (
        sector_model
        .groupby("Model")["Delta_vs_Sector_Baseline"]
        .mean()
        .sort_values(ascending=False)
    )
    best_mean_model = mean_delta.index[0]
    best_mean_delta = float(mean_delta.iloc[0])

    add_body(
        f"To test whether signal quality differs across industries, we evaluated model performance "
        f"separately by sector. Across {n_sectors} sectors, the best-performing model in each "
        f"sector beat that sector\u2019s own Always-Up baseline in {n_positive} sectors. "
        f"However, average \u0394 vs sector baseline remains negative overall "
        f"(best mean model: {best_mean_model}, \u0394={best_mean_delta:+.4f})."
    )

    p_cap3 = doc.add_paragraph(style='Normal')
    p_cap3.paragraph_format.space_after = Pt(4)
    r = p_cap3.add_run("Table 3: ")
    r.bold = True
    r.font.size = Pt(12)
    p_cap3.add_run("Best Model by Sector (Test Set)").font.size = Pt(12)

    sector_rows = []
    for _, row in sector_best.iterrows():
        sector_rows.append([
            str(row["Sector"]),
            str(row["Model"]),
            f"{int(row['N']):,}",
            f"{row['Accuracy']:.4f}",
            f"{row['Sector_Baseline']:.4f}",
            f"{row['Delta_vs_Sector_Baseline']:+.4f}",
        ])

    make_table(
        headers=["Sector", "Best Model", "N", "Accuracy", "Sector Baseline", "\u0394 vs Baseline"],
        rows=sector_rows,
        col_widths=[1.4, 1.6, 0.9, 1.1, 1.2, 1.1],
    )
else:
    add_body(
        "Per-sector result files were not found in data/model_best_by_sector.csv and "
        "data/model_sector_results.csv, so sector-level tables are omitted in this draft."
    )

add_body(
    "The gap between CV validation accuracy and test accuracy is notable for the decision tree "
    "(~56\u0025 vs 51.7\u0025), suggesting overfitting to patterns in the 2016\u20132022 "
    "training regime that do not persist in 2023\u20132026."
)

# ── Discussion ────────────────────────────────────────────────────────────
doc.add_heading('Discussion', level=2)

add_body(
    "Several factors likely contribute to the models\u2019 inability to beat the baseline:"
)

discussion_items = [
    (
        "Market efficiency.",
        " The well-documented finding that short-term price direction is difficult to predict "
        "from public information is reproduced here across all models."
    ),
    (
        "Sparse sentiment signal.",
        " News data covers only ~6\u0025 of (ticker,\u00a0date) pairs. The has_news flag "
        "and sentiment features are zero for the vast majority of rows, diluting any signal "
        "they carry."
    ),
    (
        "Survivorship bias.",
        " Using current S\u0026P\u00a0500 constituents for historical dates means the training "
        "data over-represents companies that survived and grew, potentially inflating \u201cUp\u201d "
        "rates during training."
    ),
    (
        "Regime shift.",
        " The training period (2016\u20132022) includes the COVID crash and recovery, while "
        "the test period (2023\u20132026) is a markedly different macroeconomic environment. "
        "Features calibrated on the training regime may lose predictive power."
    ),
    (
        "Class imbalance.",
        " All five models learn to predict \u201cUp\u201d most of the time. Class weights, "
        "asymmetric losses, or a tuned decision threshold could improve recall for \u201cDown\u201d days."
    ),
]

for i, (bold_part, rest) in enumerate(discussion_items, 1):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(f"{i}. ").font.size = Pt(12)
    rb = p.add_run(bold_part)
    rb.bold = True
    rb.font.size = Pt(12)
    p.add_run(rest).font.size = Pt(12)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — CONCLUSION
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Conclusion', level=1)

add_body(
    "We trained and evaluated five classification models \u2014 Decision Tree, Logistic Regression, "
    "Random Forest, XGBoost, and HistGradientBoosting \u2014 to predict next-day stock direction "
    "for S\u0026P\u00a0500 constituents "
    "using 18 features derived from price, volume, macroeconomic indicators, and news sentiment. "
    'None of the models outperformed the trivial "Always Up" baseline (52.24\u0025 accuracy) on '
    "the 2023\u20132026 test set. The highest test accuracy was 51.82\u0025 (Logistic Regression), "
    "and the highest balanced accuracy was 50.10\u0025 (HistGradientBoosting)."
)

add_body(
    "This result is consistent with the efficient market hypothesis for short-term price "
    "prediction and highlights that even a rich feature set spanning technical, macro, and "
    "sentiment signals provides minimal edge over naive rules."
)

add_body(
    "Sector-level analysis shows modest heterogeneity: a few sectors (e.g., Utilities, Energy, "
    "Consumer Cyclical, Consumer Defensive) show small positive gains versus their own sector "
    "baseline, but the average sector-level improvement remains negative across models."
)

add_body(
    "Future work could include: (1)\u00a0tuning decision thresholds and calibrating probabilities "
    "to improve minority-class recall; (2)\u00a0using class-weighted or asymmetric-loss objectives; "
    "(3)\u00a0enriching news coverage beyond the 2016\u20132020 Kaggle dataset to provide sentiment "
    "signals throughout the test period; (4)\u00a0walk-forward retraining to reduce regime drift; "
    "and (5)\u00a0narrowing the prediction task to specific sectors, regimes, or larger-move targets "
    "with potentially higher signal-to-noise."
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — REFERENCES
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('References', level=1)

references = [
    'Yang, Z. et al. (2020). "FinBERT: A Pretrained Language Model for Financial Communications." arXiv:2006.08097.',
    'Kaggle: "Massive Stock News Analysis DB for NLP/Backtests" by miguelaenlle.',
    'yfinance Python library: https://github.com/ranaroussi/yfinance',
    'Pedregosa et al. (2011). scikit-learn: Machine Learning in Python. JMLR, 12, 2825\u20132830.',
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after  = Pt(5)
    p.add_run(f"{i}. {ref}").font.size = Pt(12)

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX A — QC CHECKLIST
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Appendix A: QC Checklist', level=1)

# -- Data and claims
p_sub = doc.add_paragraph(style='Normal')
p_sub.paragraph_format.space_before = Pt(6)
p_sub.paragraph_format.space_after  = Pt(2)
rs = p_sub.add_run("Data and claims")
rs.bold = True
rs.font.size = Pt(12)

add_checklist_item("Each column\u2019s meaning is documented in scripts/merge_all.py and the EDA notebook.", checked=True)
add_checklist_item("All statistics in this report are produced by code in the notebooks.", checked=True)
add_checklist_item("Citations above should be opened and verified before final submission.", checked=False)

# -- Splits and leakage
p_sub2 = doc.add_paragraph(style='Normal')
p_sub2.paragraph_format.space_before = Pt(8)
p_sub2.paragraph_format.space_after  = Pt(2)
rs2 = p_sub2.add_run("Splits and leakage")
rs2.bold = True
rs2.font.size = Pt(12)

add_checklist_item("Time-based split at 2023-01-01 (no shuffling).", checked=True)
add_checklist_item("StandardScaler is fit on training data only.", checked=True)
add_checklist_item("Label-shuffle sanity check not yet performed \u2014 TODO before final submission.", checked=False)

# -- Models and evaluation
p_sub3 = doc.add_paragraph(style='Normal')
p_sub3.paragraph_format.space_before = Pt(8)
p_sub3.paragraph_format.space_after  = Pt(2)
rs3 = p_sub3.add_run("Models and evaluation")
rs3.bold = True
rs3.font.size = Pt(12)

add_checklist_item("Four baselines defined and reported.", checked=True)
add_checklist_item("Primary metrics are accuracy and balanced accuracy.", checked=True)
add_checklist_item("Classification reports inspected; all five models remain near-Always-Up behaviour.", checked=True)

# -- Reproducibility
p_sub4 = doc.add_paragraph(style='Normal')
p_sub4.paragraph_format.space_before = Pt(8)
p_sub4.paragraph_format.space_after  = Pt(2)
rs4 = p_sub4.add_run("Reproducibility")
rs4.bold = True
rs4.font.size = Pt(12)

add_checklist_item("README with run instructions \u2014 TODO.", checked=False)
add_checklist_item("requirements.txt / environment.yml \u2014 TODO.", checked=False)
add_checklist_item("Random seeds set: random_state=42 in tree; RANDOM_SEED=42 in baselines.", checked=True)

# ════════════════════════════════════════════════════════════════════════════
# APPENDIX B — AI USE DISCLOSURE
# ════════════════════════════════════════════════════════════════════════════

doc.add_heading('Appendix B: AI Use Disclosure', level=1)

p_tools = doc.add_paragraph(style='Normal')
p_tools.paragraph_format.space_after = Pt(4)
rb = p_tools.add_run("Tools used: ")
rb.bold = True
rb.font.size = Pt(12)
p_tools.add_run(
    "Claude Code (Anthropic, claude-sonnet-4-6), [add any others used by teammates]"
).font.size = Pt(12)

p_dates = doc.add_paragraph(style='Normal')
p_dates.paragraph_format.space_after = Pt(6)
rbd = p_dates.add_run("Dates used: ")
rbd.bold = True
rbd.font.size = Pt(12)
p_dates.add_run("[fill in]").font.size = Pt(12)

# How AI was used
p_how = doc.add_paragraph(style='Normal')
p_how.paragraph_format.space_before = Pt(4)
p_how.paragraph_format.space_after  = Pt(2)
rh = p_how.add_run("How AI was used:")
rh.bold = True
rh.font.size = Pt(12)

add_checklist_item(
    "Code generation \u2014 scripts/main.py, scripts/merge_all.py, scripts/rolling_data.py, "
    "scripts/sentiment_analysis.py, and notebook scaffolding.",
    checked=True,
)
add_checklist_item(
    "Debugging \u2014 pipeline issues with yfinance MultiIndex columns, FinBERT batching errors.",
    checked=True,
)
add_checklist_item("Writing/editing \u2014 this draft report.", checked=True)
add_checklist_item("Idea generation \u2014 [describe if applicable]", checked=False)

# What we verified
p_verified = doc.add_paragraph(style='Normal')
p_verified.paragraph_format.space_before = Pt(8)
p_verified.paragraph_format.space_after  = Pt(2)
rv = p_verified.add_run("What we verified:")
rv.bold = True
rv.font.size = Pt(12)

verified_items = [
    ("Baseline model built and compared:", " Yes (see baselines notebook)."),
    ("Leakage checks performed:", " Partial \u2014 scaler fit on train only confirmed; label-shuffle test pending."),
    ("Citations verified:", " Pending \u2014 must be completed before final submission."),
    ("Reproducibility:", " Pending \u2014 README and requirements.txt not yet written."),
]

for bold_part, rest in verified_items:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("\u2022 ").font.size = Pt(12)
    rb2 = p.add_run(bold_part)
    rb2.bold = True
    rb2.font.size = Pt(12)
    p.add_run(rest).font.size = Pt(12)

# AI mistake example
p_mistake = doc.add_paragraph(style='Normal')
p_mistake.paragraph_format.space_before = Pt(8)
p_mistake.paragraph_format.space_after  = Pt(2)
rm = p_mistake.add_run("One example of an AI mistake we caught:")
rm.bold = True
rm.font.size = Pt(12)

mistake_items = [
    ("What AI suggested:", " [fill in a real example from your workflow]"),
    ("Why it was wrong / risky:", " [explain]"),
    ("What we did instead:", " [explain]"),
]

for bold_part, rest in mistake_items:
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("\u2022 ").font.size = Pt(12)
    rb3 = p.add_run(bold_part)
    rb3.bold = True
    rb3.font.size = Pt(12)
    p.add_run(rest).font.size = Pt(12)

# ════════════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
